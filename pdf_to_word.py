#!/usr/bin/env python3
"""
Konvertira PDF s tekstovima pjesama u Word dokument.
Svaka stranica PDF-a = jedna pjesma.

Korištenje:
    python pdf_to_word.py ulaz.pdf izlaz.docx
    python pdf_to_word.py ulaz.pdf          # sprema kao ulaz.docx
"""

import sys
import re
from pathlib import Path

import pdfplumber
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Konstante za formatiranje ──────────────────────────────────────────────────
FONT_NAME = "Liberation Serif"
FONT_SIZE_TITLE = 26
FONT_SIZE_NORMAL = 9
COLOR_BLACK = RGBColor(0, 0, 0)

# Fontovi s punom podrškom za hrvatska slova (isti izbor kao original)
# Liberation Serif ≈ Times New Roman | Liberation Sans ≈ Arial
FONT_CHOICES = ["Liberation Serif", "Liberation Sans", "FreeSerif", "FreeSans"]

# Zadani stil; web sučelje šalje vlastite vrijednosti
STYLE_DEFAULTS = {
    "layout": "kompaktno",       # "kompaktno" (stupci) ili "klasicno" (stranice)
    "font": FONT_NAME,
    "title_size": FONT_SIZE_TITLE,
    "body_size": FONT_SIZE_NORMAL,
    "margin_cm": 0.5,
    "page_break": True,
    "n_cols": 3,
    "strip_chords": True,
}

# Labele sekcija koje se ispisuju boldano
SECTION_LABELS = re.compile(
    r"^(Chorus\s*\d*|Verse\s*\d*|Bridge\s*\d*|Intro|Outro|Pre-?Chorus|Tag|"
    r"V\d+(?:\s*\([^)]+\))?|C\d*|B\d*|BRIDGE|Verse|verse|chorus)\s*(\(.*\))?\s*$",
    re.IGNORECASE,
)

# Sekcije koje se u kompaktnom izgledu tretiraju kao REFREN (uvučen, bold+kurziv)
CHORUS_TYPES = re.compile(
    r"^(Chorus\s*\d*|C\d*|Bridge\s*\d*|B\d*|BRIDGE|Pre-?Chorus|Tag|Outro)\s*(\(.*\))?$",
    re.IGNORECASE,
)

# Crta razdjelnica između pjesama u kompaktnom izgledu
DIVIDER = "_" * 35

# Regex za jedan akord (hrvatska/njemačka notacija)
# Primjeri: D, A, Fis, H, Cis, Es, Dm, Am7, G/B, Dsus4, Hm, Bb
_CHORD_TOKEN = re.compile(
    r"^[CDEFGAHB]"          # osnovna nota
    r"(is|es|s|#|b)?"       # povisilica/snizilica (Fis, Es, As, C#, Bb)
    r"(m|mol|maj|min|dim|aug|sus|add)?"  # kvaliteta
    r"\d*"                  # broj (7, 9, 11...)
    r"(sus\d*|add\d*)?"     # sus/add sufiks
    r"(/[CDEFGAHB](IS|ES|is|es|#|b)?)?$",  # slash akord
    re.IGNORECASE,
)


def is_chord_line(line: str) -> bool:
    """Vraća True ako linija sadrži samo akorde (bez normalnih riječi)."""
    tokens = line.split()
    if not tokens:
        return False
    return all(bool(_CHORD_TOKEN.match(t)) and len(t) <= 7 for t in tokens)


def _style(style=None):
    merged = dict(STYLE_DEFAULTS)
    if style:
        merged.update({k: v for k, v in style.items() if v is not None})
    return merged


# ── Pomoćne funkcije za Word ───────────────────────────────────────────────────

def set_paragraph_spacing(para, before=0, after=0, line_spacing=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line_spacing is not None:
        spacing.set(qn("w:line"), str(line_spacing))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def add_run(para, text, bold=False, size=FONT_SIZE_NORMAL, italic=False,
            font=FONT_NAME):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = COLOR_BLACK
    return run


def add_title_paragraph(doc, title_text, key_text, st):
    """Naslov lijevo + Key desno u istom odlomku s tabulatorom."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(para, before=0, after=60)

    # Postavi tab stop na desni rub stranice
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9070")  # ~16 cm (širina teksta)
    tabs.append(tab)
    pPr.append(tabs)

    run_title = para.add_run(title_text)
    run_title.font.name = st["font"]
    run_title.font.size = Pt(st["title_size"])
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_BLACK

    if key_text:
        run_tab = para.add_run("\t")
        run_tab.font.name = st["font"]
        run_tab.font.size = Pt(st["body_size"])

        run_key = para.add_run(key_text)
        run_key.font.name = st["font"]
        run_key.font.size = Pt(st["body_size"])
        run_key.font.bold = False
        run_key.font.color.rgb = COLOR_BLACK

    return para


def add_section_label(doc, label, st):
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before=120, after=0)
    add_run(para, label, bold=True, size=st["body_size"], font=st["font"])
    return para


def add_lyric_line(doc, text, st):
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before=0, after=0)
    add_run(para, text, size=st["body_size"], font=st["font"])
    return para


def add_meta_label(doc, label, values, st):
    """'Pjesmarica' bold, ispod njega vrijednosti normalno."""
    para_label = doc.add_paragraph()
    set_paragraph_spacing(para_label, before=0, after=0)
    add_run(para_label, label, bold=True, size=st["body_size"], font=st["font"])

    for val in values:
        para_val = doc.add_paragraph()
        set_paragraph_spacing(para_val, before=0, after=0)
        add_run(para_val, val, size=st["body_size"], font=st["font"])

    return para_label


def add_empty_line(doc, st):
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before=0, after=0)
    add_run(para, "", size=st["body_size"], font=st["font"])
    return para


# ── Parsiranje jedne stranice PDF-a ───────────────────────────────────────────

def parse_page(page):
    """
    Vraća dict:
        title      : str
        key        : str  (npr. "Key: A (Ab)\nCapo 1")
        subtitle   : str | None  (redak ispod naslova, lijevo)
        pjesmarica : list[str]
        sections   : list of (label, lines)
            label = str ili None (za blokove bez labele)
            lines = list[str]
        notes      : list[str]   (redovi PRIJE prve sekcije, a iza Pjesmarica)
    """
    # Dohvati sve riječi s pozicijama
    words = page.extract_words(x_tolerance=3, y_tolerance=3)
    if not words:
        return None

    page_width = page.width
    right_threshold = page_width * 0.55  # desna polovina = Key zona

    # Grupiraj u retke po y koordinati
    lines_by_y = {}
    for w in words:
        y = round(w["top"] / 3) * 3  # zaokruži na 3pt
        lines_by_y.setdefault(y, []).append(w)

    sorted_ys = sorted(lines_by_y.keys())
    raw_lines = []  # lista (y, left_x, text_left, text_right)
    for y in sorted_ys:
        row_words = sorted(lines_by_y[y], key=lambda w: w["x0"])
        left_words = [w for w in row_words if w["x0"] < right_threshold]
        right_words = [w for w in row_words if w["x0"] >= right_threshold]
        text_left = " ".join(w["text"] for w in left_words).strip()
        text_right = " ".join(w["text"] for w in right_words).strip()
        left_x = row_words[0]["x0"] if row_words else 0
        raw_lines.append((y, left_x, text_left, text_right))

    if not raw_lines:
        return None

    # ── 1. Naslov: prva linija (najveći font, bold) ──────────────────────────
    # Detektiramo naslov kao prvu nepraznu lijevu liniju
    title = ""
    subtitle = None
    key_parts = []

    title_idx = 0
    for i, (y, lx, tl, tr) in enumerate(raw_lines):
        if tl:
            title = tl
            title_idx = i
            # Key može biti na istom retku (desno)
            if tr:
                key_parts.append(tr)
            break

    # Redci odmah ispod naslova sa SAMO desnim tekstom = nastavak Key
    # (npr. "Key: A" pa "Capo 1" u zasebnim redovima)
    next_idx = title_idx + 1
    while next_idx < len(raw_lines):
        _, _, tl2, tr2 = raw_lines[next_idx]
        if tr2 and not tl2:
            key_parts.append(tr2)
            next_idx += 1
        else:
            break

    # Redak ispod – može biti podnaslov (italic/manji font) ili "No" itd.
    # Heuristika: ako sljedeći red ima tekst i nije "Pjesmarica", tretiramo
    # ga kao podnaslov ako ne počinje s "Key:" i nije broj
    if next_idx < len(raw_lines):
        y2, lx2, tl2, tr2 = raw_lines[next_idx]
        if tl2 and not tl2.startswith("Pjesmarica") and not re.match(r"^\d+$", tl2):
            # kratki redak bez labels = podnaslov
            if len(tl2.split()) <= 6 and not SECTION_LABELS.match(tl2):
                # provjeri je li sljedeći redak Pjesmarica ili prazan
                peek_idx = next_idx + 1
                if peek_idx < len(raw_lines):
                    _, _, tl3, _ = raw_lines[peek_idx]
                    if tl3.startswith("Pjesmarica") or not tl3:
                        subtitle = tl2
                        next_idx += 1
                        # provjeri Capo na desnoj strani tog retka
                        if tr2:
                            key_parts.append(tr2)

    key_text = "\n".join(key_parts) if key_parts else ""

    # ── 2. Pjesmarica blok ──────────────────────────────────────────────────
    pjesmarica_values = []
    pi = next_idx
    if pi < len(raw_lines) and raw_lines[pi][2].startswith("Pjesmarica"):
        pi += 1
        while pi < len(raw_lines):
            _, _, tl, _ = raw_lines[pi]
            if not tl or SECTION_LABELS.match(tl) or tl.startswith("Verse") \
               or tl.startswith("Chorus") or tl.startswith("Bridge"):
                break
            if re.match(r"^[\d\s\-/]+$", tl) or re.match(r"^P\d", tl):
                pjesmarica_values.append(tl)
                pi += 1
            else:
                break

    # ── 3. Ostatak = napomene + sekcije ─────────────────────────────────────
    notes = []
    sections = []  # list of [label, lines_list]
    current_label = None
    current_lines = []
    in_notes = True  # dok ne dođemo do prve sekcije

    i = pi
    while i < len(raw_lines):
        y, lx, tl, tr = raw_lines[i]
        i += 1

        if not tl:
            # prazan red
            if current_label is not None or current_lines:
                current_lines.append("")
            elif in_notes and notes:
                notes.append("")
            continue

        if SECTION_LABELS.match(tl):
            # spremi prethodni blok
            if in_notes and current_lines:
                notes.extend(current_lines)
                current_lines = []
            elif current_label is not None or current_lines:
                # trim trailing empty
                while current_lines and current_lines[-1] == "":
                    current_lines.pop()
                sections.append([current_label, current_lines])
                current_lines = []

            current_label = tl
            in_notes = False
        else:
            if in_notes:
                notes.append(tl)
            else:
                current_lines.append(tl)

    # spremi zadnji blok
    if current_label is not None or current_lines:
        while current_lines and current_lines[-1] == "":
            current_lines.pop()
        sections.append([current_label, current_lines])

    # trim notes trailing empty
    while notes and notes[-1] == "":
        notes.pop()

    return {
        "title": title,
        "key": key_text,
        "subtitle": subtitle,
        "pjesmarica": pjesmarica_values,
        "notes": notes,
        "sections": sections,
    }


def parse_pdf(pdf_path, log=None):
    """Parsira cijeli PDF -> lista pjesama (dict po parse_page)."""
    songs = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            song = parse_page(page)
            if song and song["title"]:
                songs.append(song)
                if log:
                    log(f"  Stranica {page_num}: {song['title']}")
            elif log:
                log(f"  Stranica {page_num}: (preskočena – nema sadržaja)")
    return songs


# ── Pisanje jedne pjesme u Word dokument ──────────────────────────────────────

def write_song(doc, song, add_page_break=False, style=None):
    st = _style(style)
    if add_page_break:
        doc.add_page_break()

    # Naslov + Key
    add_title_paragraph(doc, song["title"], song["key"], st)

    # Podnaslov (npr. "Božja Pobjeda", "Papa Band")
    if song["subtitle"]:
        para = doc.add_paragraph()
        set_paragraph_spacing(para, before=0, after=0)
        add_run(para, song["subtitle"], size=st["body_size"], font=st["font"])

    # Pjesmarica
    if song["pjesmarica"]:
        add_empty_line(doc, st)
        add_meta_label(doc, "Pjesmarica", song["pjesmarica"], st)

    # Napomene (tekst prije sekcija)
    if song["notes"]:
        add_empty_line(doc, st)
        for note in song["notes"]:
            if note == "":
                add_empty_line(doc, st)
            else:
                add_lyric_line(doc, note, st)

    # Sekcije
    for label, lines in song["sections"]:
        add_empty_line(doc, st)
        if label:
            add_section_label(doc, label, st)
        # Grupiraj linije u blokove odvojene praznim recima
        for line in lines:
            if line == "":
                add_empty_line(doc, st)
            else:
                add_lyric_line(doc, line, st)


# ── Postavljanje margina dokumenta ────────────────────────────────────────────

def setup_document(style=None):
    st = _style(style)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(st["margin_cm"])
    section.bottom_margin = Cm(st["margin_cm"])
    section.left_margin = Cm(st["margin_cm"])
    section.right_margin = Cm(st["margin_cm"])

    # Ukloni zadani razmak između odlomaka
    doc_style = doc.styles["Normal"]
    doc_style.font.name = st["font"]
    doc_style.font.size = Pt(st["body_size"])
    pf = doc_style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    return doc


def build_docx(songs, style=None):
    """Složi Word dokument iz liste pjesama; vraća Document objekt."""
    st = _style(style)
    doc = setup_document(st)
    for idx, song in enumerate(songs):
        add_break = idx > 0 and st["page_break"]
        if idx > 0 and not st["page_break"]:
            add_empty_line(doc, st)
            add_empty_line(doc, st)
        write_song(doc, song, add_page_break=add_break, style=st)
    return doc


# ── Kompaktni izgled (A4 ležeće, stupci, bez naslova — original konverter) ────

CHORUS_INDENT = Twips(720)


def set_columns(doc, n_cols, margin_cm=0.5):
    """A4 ležeće s n_cols stupaca i minimalnim marginama."""
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)

    sectPr = section._sectPr
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(n_cols))
    cols.set(qn("w:space"), "720")
    cols.set(qn("w:equalWidth"), "1")
    sectPr.append(cols)


def song_compact_blocks(song):
    """Pjesma -> [(je_refren, linija), ...] za kompaktni ispis.

    Naslov, tonalitet i Pjesmarica blok se izostavljaju; napomene i sekcije
    bez labele idu kao kitica (normalno), Chorus/Bridge/Tag/Outro kao refren
    (uvučeno, bold+kurziv). Prazni retci se preskaču.
    """
    out = []
    for note in song.get("notes") or []:
        if note:
            out.append((False, note))
    for label, lines in song.get("sections") or []:
        chorus = bool(label and CHORUS_TYPES.match(label))
        for line in lines:
            if line:
                out.append((chorus, line))
    return out


def build_docx_compact(songs, style=None):
    """Kompaktni Word: A4 ležeće, stupci, crta između pjesama."""
    st = _style(style)
    doc = Document()

    doc_style = doc.styles["Normal"]
    doc_style.font.name = st["font"]
    doc_style.font.size = Pt(st["body_size"])
    pf = doc_style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    set_columns(doc, st["n_cols"], st["margin_cm"])

    first = True
    for song in songs:
        if not first:
            para = doc.add_paragraph()
            set_paragraph_spacing(para, before=60, after=60)
            add_run(para, DIVIDER, size=st["body_size"], font=st["font"])
        first = False
        for chorus, line in song_compact_blocks(song):
            para = doc.add_paragraph()
            set_paragraph_spacing(para, before=0, after=0)
            if chorus:
                para.paragraph_format.left_indent = CHORUS_INDENT
            add_run(para, line, bold=chorus, italic=chorus,
                    size=st["body_size"], font=st["font"])
    return doc


# ── Glavna funkcija ────────────────────────────────────────────────────────────

def convert_pdf_to_docx(pdf_path: str, docx_path: str):
    print(f"Učitavam: {pdf_path}")
    songs = parse_pdf(pdf_path, log=print)

    if not songs:
        print("Nije pronađena nijedna pjesma!")
        return

    doc = build_docx(songs)
    doc.save(docx_path)
    print(f"\nSpremi kao: {docx_path}")
    print(f"Ukupno pjesama: {len(songs)}")


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    pdf_in = sys.argv[1]
    if len(sys.argv) >= 3:
        docx_out = sys.argv[2]
    else:
        docx_out = str(Path(pdf_in).with_suffix(".docx"))

    convert_pdf_to_docx(pdf_in, docx_out)
